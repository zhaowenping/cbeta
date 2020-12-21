#!/usr/bin/env python
# -*- coding: UTF-8 -*-
# Language Version: 2.7+
# Last Modified: 2020-12-20 19:27:22
from __future__ import unicode_literals, division, absolute_import, print_function

"""
函数库
1. 简体繁体检测
2. 简体繁体转换
3. 注音 phonetic notation
4. 注音格式转换
5. 去除'〡'
6. 所有数据装入redis
"""

__all__ = []
__author__ = ""
__version__ = "0.0.1"


import re
import os
import copy
import gzip
import json
import time
import array
from functools import reduce
from functools import total_ordering
import unicodedata
# import xml.dom.minidom as minidom
from xml.etree import ElementTree as ET
import pprint

# docx
import requests
from docx import Document
from docx.shared import Inches, Emu
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt
from docx.shared import RGBColor
import Levenshtein


print('调用函数库')
PATH = "/home/zhaowp/cbeta/cbeta"
# 拉丁字母0041~024F


def VERTICAL(ctx):
    '''竖排标点转横排标点'''
    t = { 0xfe10: 0xff0c,  # 逗号
          0xfe11: 0x3001,  # 顿号
          0xfe12: 0x3002,  # 句号
          0xfe13: 0xff1a,  # 冒号
          0xfe14: 0xff1b,  # 分号
          0xfe15: 0xff01,  # 感叹号
          0xfe16: 0xff1f,  # 问号
          0xfe17: 0x3016,  # 左括号
          0xfe18: 0x3017,  # 右括号
          0xfe19: 0x2026,  # 省略号
          0xfe31: 0x2014,  # 破折号
          0xfe35: 0xff08,  # 左小括号
          0xfe36: 0xff09,  # 右小括号
          0xfe41: 0x300c,  # 左内引号
          0xfe42: 0x300d,  # 右内引号
          0xfe43: 0x300e,  # 左外引号
          0xfe44: 0x300f,  # 右外引号
            }
    # ！＂＃＄％＆＇（）＊＋，－．／
    # ：；＜＝＞？＠
    # ［＼］＿
    # ｛｜｝～
    ctx = ctx.translate(t)
    return ctx


def fullwidth2half(ctx, pun=True, exp=''):
    '''全角字符转半角字符, pun: 默认转换全角标点, exp: 不转换的字符列表'''

    if pun:
        # '！' <= ch <= '～':  # 0xFF01~0xFF5E
        t = {ch: ch-0xFEE0 for ch in range(0xFF01, 0xFF5F)}
        t[0x3000] = 0x20
    else:
        # if '０' <= ch <= '９' or 'Ａ' <= ch <= 'Ｚ' or 'ａ' <= ch <= 'ｚ':
        t = {ch: ch-0xFEE0 for ch in range(0xFF10, 0xFF1A)}
        t.update({ch: ch-0xFEE0 for ch in range(0xFF21, 0xFF3B)})
        t.update({ch: ch-0xFEE0 for ch in range(0xFF41, 0xFF5B)})
    for ch in exp:
        t.pop(ord(ch), None)
    ctx = ctx.translate(t)
    return ctx

    # ctx = array.array('u', ctx)
    # SPACE = chr(0x3000)
    # if pun:
    #     for idx, ch in enumerate(ctx):
    #         if ch == SPACE:
    #             ctx[idx] = ' '
    #         elif '！' <= ch <= '～':  # 0xFF01~0xFF5E
    #             ctx[idx] = chr(ord(ch) - 0xFEE0)
    # else:
    #     for idx, ch in enumerate(ctx):
    #         if '０' <= ch <= '９' or 'Ａ' <= ch <= 'Ｚ' or 'ａ' <= ch <= 'ｚ':
    #             ctx[idx] = chr(ord(ch) - 0xFEE0)

    # return ctx.tounicode()


def python_unescape(ctx):
    '''替换python字样转义字符串为正常汉字'''
    for ch in re.findall(r'(?:[^\\]|^)(\\u[a-fA-F0-9]{4})', ctx):
        ctx = ctx.replace(ch, chr(int(ch[-4:], 16)))
    for ch in re.findall(r'(?:[^\\]|^)(\\U[a-fA-F0-9]{8})', ctx):
        ctx = ctx.replace(ch, chr(int(ch[-8:], 16)))
    # ctx = ctx.replace(r'\\\\', r'\\')

    return ctx


def python_escape(ctx):
    '''因为ES不支持F区、G区汉字，所以将F区、G区汉字转换成转义字符序列，方便后续ES中查找'''
    for char in ctx:
        if 0x2CEB0 <= ord(char) <= 0x2EBE0:  # F区
            yield r'\U{:08X}'.format(ord(char))
        elif 0x30000 <= ord(char) <= 0x3134A:  # G区
            yield r'\U{:08X}'.format(ord(char))
        else:
            yield char


def ids_split(ctx, fn=lambda x:x, exfn=lambda x:x):
    '''将字符串按照ids序列分割为数组。遇到ids序列的时候，使用fn处理，返回处理后的数组。'''
    stack = []
    counter = 0
    for ch in ctx:
        if counter == 0:
            if stack and ch in '↷↹⿰⿱⿴⿵⿶⿷⿸⿹⿺⿻⿲⿳':
                yield exfn(''.join(stack))
                stack = []
            stack.append(ch)
            if ch in '↷↹':
                counter = 1
            elif ch in '⿰⿱⿴⿵⿶⿷⿸⿹⿺⿻':
                counter = 2
            elif ch in '⿲⿳':
                counter = 3
        else:
            stack.append(ch)
            if ch in '↷↹':
                pass
            elif ch in '⿰⿱⿴⿵⿶⿷⿸⿹⿺⿻':
                counter += 1
            elif ch in '⿲⿳':
                counter += 2
            else:
                counter -= 1
                if counter == 0:
                    yield fn(''.join(stack))
                    stack = []
    yield exfn(''.join(stack))


def hz_len(ctx):
    '''判断一个unicode文本的汉字长度: 有几个汉字组成'''
    return sum(ids_split(ctx, fn=lambda x:1, exfn=len))


class IDS:
    def __init__(self):
        self.ids_dict = dict()
        with open('idx/ids.txt') as fd:
            for line in fd:
                line = line.strip().split()
                self.ids_dict[line[2]] = line[1]

    def rm_ids(self, ctx):
        '''替换unicode ids形式为普通单个字符, 无法替换则保持原样不动'''
        # return ''.join(ids_split(ctx, fn=lambda x:self.ids_dict.get(x, x)))
        return ''.join(ids_split(ctx, fn=self.find_ids))

    def find_ids(self, ctx):
        '''查找ids序列
        牛 on left = 牜
        玉 on left = 𤣩
        示 on left = 礻
        竹 on top = 𥫗
        糸 on left = 糹
        肉 on left = 月
        艸 on top = 艹
        言 on left = 訁
        金 on left = 釒
        食 on left = 飠
        牜𤣩礻糹⺼艹訁釒飠氵冫忄
        '''
        tt = {'牛': '牜', '王': '𤣩', '示': '礻', '糸': '糹', '月': '⺼', '草': '艹', '言': '訁',
              '金': '釒', '食': '飠', '水': '氵', '冰': '冫', '心': '忄', '人': '亻', '衣': '衤',
              '手': '扌', '犬': '犭', '病': '疒', '爪': '爫', '火': '灬', '足': '𧾷'}
        rr = self.ids_dict.get(ctx, None)
        if rr:
            return rr
        # 没找到就替换之后再找一次
        tt = {ord(k): ord(tt[k]) for k in tt}
        ctx = ctx.translate(tt)  # .replace(chr(0xFFFD), '')
        rr = self.ids_dict.get(ctx, None)
        if rr:
            return rr
        return ctx


class CBETA_COM:
    def __init__(self):
        with open('idx/composition.json') as fd:
            self.desc = json.load(fd)

    def rm_com(self, ctx):
        ''' 替换cbeta组字式为相应的字符, 无法替换则替换为空格'''
        return ''.join(re_split(r'\[.*?\]', ctx, lambda x: self.desc.get(x, ' ')))


def rm_joiner(ctx):
    '''去除汉字的链接和装饰符号: 外圈加方框、圆形或者三角'''
    tt = {0x200C:0xFFFD, 0x200D:0xFFFD, 0x20DD:0xFFFD, 0x20DE:0xFFFD, 0x20DF:0xFFFD,
            0x20E0:0xFFFD, 0x20E4:0xFFFD, }
    ctx = ctx.translate(tt).replace(chr(0xFFFD), '')
    return ctx


# _space = re.compile(r'[ \t\n\r\x0b\x0c\u3000]+')
def normalize_space(ctx):
    '''去掉字符串两边的空格, 中间连续的多个空格(空白)替换为一个'''
    ctx = ctx.strip()
    ctx = re.sub(r'\s+', ' ', ctx)
    return ctx


def rm_ditto_mark(ctx):
    # 在xml中去除四个叠字符号(默认叠字符号始终相连): ⺀ U+2E80 0 /〃 U+3003 2227 /々 U+3005 6415/ 亽 U+4EBD 151/ 𠚤 U+206A4
    # 首先删除叠字符号中间的空白
    ctx = re.sub(r'([\u3003\u3005\u4ebd\U000206A4])\s+([\u3003\u3005\u4ebd\U000206A4])', r'\1\2', ctx)
    ctx = array.array('u', ctx)
    dittos = (chr(0x3003), chr(0x3005), chr(0x4ebd), chr(0x206A4))
    cc = 0  # 叠字符号的重复次数
    len_ctx = len(ctx)
    for idx, zi in enumerate(ctx):
        if zi in dittos:
            cc = cc + 1
            if len_ctx > idx+1 and ctx[idx+1] in dittos:
                continue
        if cc == 0:
            continue
        j = 0
        for i in range(idx-cc, -1, -1):
            if ishanzi(ctx[i]):
                ctx[idx-j] = ctx[i]  # 找到一个合法的重复字符进行替换
                j = j + 1
                cc = cc - 1
                if cc == 0:
                    break
    return ctx.tounicode()


def ishanzi(zi):
    '''粗略判断一个字符是否是非叠字汉字'''
    zi = ord(zi)
    if zi in {'\u2E80', '\u3003', '\u3005', '\u4ebd', '\U000206A4'}:
        return False
    # 主区和A区, 包括补充区0x4DB6-0x4DBF, 0x9FA6-0x9FFC
    if 0x3400 <= zi <= 0x9FFC: # and zi != 0x4EBD:
        return True
    # BCDEFG，包括兼容汉字区0x2FXXX
    if 0x20000 <= zi <= 0x3134A:
        return True
    # 〇
    if 0x3007 == zi:
        return True
    # 兼容汉字区
    if 0xF900 <= zi <= 0xFAD9:
        return True
    return False


def unicode_zone(char):
    '''汉字编码范围'''
    if 0x4E00 <= ord(char) <= 0x9FA5:
        return 'M'
    if 0x9FA6 <= ord(char) <= 0x9FFC:
        return 'Mx'
    if 0x3400 <= ord(char) <= 0x4DB5:
        return 'A'
    if 0x4DB6 <= ord(char) <= 0x4DBF:
        return 'Ax'
    if 0x20000 <= ord(char) <= 0x2A6D6:
        return 'B'
    if 0x2A6D7 <= ord(char) <= 0x2A6DD:
        return 'Bx'
    if 0x2A700 <= ord(char) <= 0x2B734:
        return 'C'
    if 0x2B740 <= ord(char) <= 0x2B81D:
        return 'D'
    if 0x2B820 <= ord(char) <= 0x2CEA1:
        return 'E'
    if 0x2CEB0 <= ord(char) <= 0x2EBE0:
        return 'F'
    if 0x30000 <= ord(char) <= 0x3134A:
        return 'G'
    return ''


def readdb(path, trans=False, reverse=False):
    '''读取文本数据库, trans为是否用于tanslate函数, reverse为是否翻转'''
    result = dict()
    #path = os.path.join("/home/zhaowp/cbeta/cbeta", path)
    with open(path, encoding='utf8') as fd:
        for line in fd:
            line = line.strip()
            if line.startswith('#') or not line: continue
            c0, c1, *cc = line.strip().split()
            if trans and reverse:
                result[ord(c1)] = ord(c0)
            if trans and not reverse:
                result[ord(c0)] = ord(c1)
            if not trans and reverse:
                result[c1] = c0
            if not trans and not reverse:
                result[c0] = c1
    return result



def read_menu_file(sutra_list):
    '''读取tab分隔的菜单文件，返回树状字典'''
    menu = dict()
    print('''读取tab分隔的菜单文件，返回树状字典''')

    with open(sutra_list, encoding='utf8') as fd:
        for line in fd:
            line = line.rstrip()
            # if line.startswith('\t\t\t\t\t'):
            #     print(line)
            if not line.startswith('\t'):
                key1 = line
                menu.update({line:{}})
                continue
            line = line[1:]

            if not line.startswith('\t'):
                key2 = line
                menu[key1].update({line: {}})
                continue
            line = line[1:]

            if not line.startswith('\t'):
                key3 = line
                menu[key1][key2].update({line: {}})
                continue
            line = line[1:]

            if not line.startswith('\t'):
                key4 = line
                menu[key1][key2][key3].update({line: {}})
                continue
            line = line[1:]

            if not line.startswith('\t'):
                key5 = line
                menu[key1][key2][key3][key4].update({line: {}})
                continue
            line = line[1:]

            if not line.startswith('\t'):
                menu[key1][key2][key3][key4][key5].update({line: {}})
                continue
        return menu

# def ls(pattern):
#     result = []
#     for path in os.listdir(pattern):
#         if path.startswith(sutra) and re.match(pattern, path):
#             result.append(path.split('_')[1][:-4])
#     return result


def normalize_text(ctx):
    '''标准化文本(只适合繁体字)'''
    # # 去除上标和下标
    # ctx = re.sub(r'\[[\w\*]+\]', '', ctx)
    # # 去除组字式
    # ctx = rm_com(ctx)
    # 去除错误的标点符号
    tt = {0xff0e: 0x00b7,
          0x2027: 0x00b7,
          0x25CB: ord('〇'),  # 佛光山大辞典的用法
          }
    ctx = ctx.translate(tt)
    # 去除两边空格及多余空格
    ctx = normalize_space(ctx)
    # 去除汉字链接符号
    # ctx = rm_joiner(ctx)
    # 去除汉字重复符号
    ctx = rm_ditto_mark(ctx)
    # 去除异体字、异体词
    ctx = rm_variant(ctx)
    return ctx

def get_first_juan(number):
    '''给定经号T01n0002,T20n1113B, T03n0154_002#p0085a13,  返回所有排序后的卷['001', '002', ...]中的第一卷
    返回值是一个数字，如果没有找到则返回0'''
    if '_' in number:
        number = number.split('_')[0]
    book, sutra = number.split('n')
    # 查找第一卷(有些不是从第一卷开始的)
    juan = 999
    if not os.path.exists(f'xml/{book}'):
        return 0
    for path in os.listdir(f'xml/{book}'):
        if path.startswith(number):
            juan = min(juan, int(path.split('_')[1][:3]))
    if juan == 999:
        juan = 0
    return juan

# lb_anchor 0607a18  #p0481a25
# pb_anchor T02.0125.0603c
@total_ordering
class Number:
    '''经号类: T01n0002a_002'''
    def __init__(self, n, anchor=''):
        if isinstance(n, tuple):
            self.book, tome, self.sutra, self.yiyi, self.volume, self.anchor = n
        else:
            self.book, tome, self.sutra, self.yiyi, self.volume, self.anchor = None, '', '', '', 0, ''
            # r = re.findall(r'([A-Z]{1,2})(\d{2,3})n(\w\d{3})([a-zA-Z])?(?:_(\d{3}))?', n)
            r = re.findall(r'([A-Z]{1,2})(\d{2,3})n(\w\d{3})([a-zA-Z])?(?:_(\d{3}))?(?:#p(\w\d{3}[abc]\d\d))?', n)
            if r:
                self.book, tome, self.sutra, self.yiyi, volume, self.anchor = r[0]
                self.volume = 0 if not volume else int(volume)
        # 重新标准化tome
        tome = int(tome)
        tome_len = 2
        if self.book in {'A', 'C', 'G', 'GA', 'GB', 'L', 'M', 'P', 'U'}:
            tome_len = 3
        self.tome = f'{tome:0{tome_len}}'

    def __eq__(self, other):
        return (self.book == other.book and
                self.tome == other.tome and
                self.sutra == other.sutra and
                self.yiyi == other.yiyi and
                self.volume == other.volume)

    def __lt__(self, other):
        '''重要性的排序'''
        pr = ("T", "A", "S", "F", "C", "K", "B", "ZW", "P", "U", "D", "G", "M", "N", "L", "J", "X", "Y", "LC", "GA", "GB", "I")
        tt = {  'T': 1, 'A': 2, 'F': 3, 'S': 4, 'U': 5,
                'P': 6, 'B': 7, 'ZW': 8, 'J': 9, 'C': 10,
                'D': 11, 'K': 12, 'G': 13, 'N': 18, 'I':19,
                'L': 20, 'M': 30, 'Y':40, 'LC':50,
                'GA':60, 'GB': 70, 'ZS':80, 'X':90}
        yiyi = 0 if not self.yiyi else int(self.yiyi, 16)
        oyiyi = 0 if not other.yiyi else int(other.yiyi, 16)
        return (tt[self.book], int(self.tome), int(self.sutra, 16), yiyi, self.volume) < (tt[other.book], int(other.tome), int(other.sutra, 16), oyiyi, other.volume)

    def __str__(self):
        if not self.book:
            return ''
        if self.volume:
            return f'{self.book}{self.tome}n{self.sutra}{self.yiyi}_{self.volume:03}'
        else:
            return f'{self.book}{self.tome}n{self.sutra}{self.yiyi}'

    @property
    def url(self):
        if self.volume:
            volume = self.volume
        else:
            volume = self.get_first_juan()
        if self.anchor:
            return f'/xml/{self.book}{self.tome}/{self.book}{self.tome}n{self.sutra}{self.yiyi}_{volume:03}.xml#{self.anchor}'
        else:
            return f'/xml/{self.book}{self.tome}/{self.book}{self.tome}n{self.sutra}{self.yiyi}_{volume:03}.xml'

    @property
    def pages(self):
        '''大般若经等....'''
        if f'{self.book}{self.sutra}' == 'T0220':
            return list(range(1, 601))
        number = f'{self.book}{self.tome}n{self.sutra}{self.yiyi}'
        if not os.path.exists(f'xml/{self.book}{self.tome}'):
            return []
        pages = [int(path.split('_')[1][:3]) for path in os.listdir(f'xml/{self.book}{self.tome}') if path.startswith(number)]
        return sorted(pages)

    def get_first_juan(self):
        '''给定经号T01n0002，返回所有排序后的卷['001', '002', ...]中的第一个
        返回值是一个数字，如果没有找到则返回0'''
        number = f'{self.book}{self.tome}n{self.sutra}{self.yiyi}'
        # 查找第一卷(有些不是从第一卷开始的)
        juan = 999
        if not os.path.exists(f'xml/{self.book}{self.tome}'):
            return 0
        for path in os.listdir(f'xml/{self.book}{self.tome}'):
            if path.startswith(number):
                juan = min(juan, int(path.split('_')[1][:3]))
        if juan == 999:
            juan = 0
        return juan

    def __add__(self, page):
        '''给定经号T01n0002_001，返回T01n0002_002'''
        # 重新生成标准经号
        if not self.volume:
            pass  #raise 一个错误?
        number = f'{self.book}{self.tome}n{self.sutra}{self.yiyi}_{self.volume:03}'
        # # 对所有的book下的卷排序
        juanlist = get_sorted_juan(f'{self.book}{self.tome}')
        page = juanlist.index(number) + page
        if 0 <= page < len(juanlist):
            return Number(juanlist[page])
        # else: tome + 1
        # tomelist: 获得全部book(T01)下的所有排序好的册号列表(T01,T02,T03,T04...)
        page = page - len(juanlist)
        tomelist = (path.strip(self.book) for path in os.listdir(f'xml') if path.startswith(self.book))
        tomelist = tuple(f'{self.book}{i}' for i in sorted(tomelist, key=int))
        idx = tomelist.index(f'{self.book}{self.tome}')
        if idx + 1 < len(tomelist):
            nextbook = tomelist[idx + 1]
            juanlist = get_sorted_juan(nextbook)
            return Number(juanlist[page])
        # else: book + 1
        return juanlist


    def __sub__(self, page):
        '''给定经号T01n0002_002，返回T01n0002_001'''
        # 重新生成标准经号
        # print('page:', page)
        if not self.volume:
            pass  #raise 一个错误?
        number = f'{self.book}{self.tome}n{self.sutra}{self.yiyi}_{self.volume:03}'
        # # 对所有的book下的卷排序
        juanlist = get_sorted_juan(f'{self.book}{self.tome}')
        page = juanlist.index(number) - page
        if 0 <= page < len(juanlist):
            return Number(juanlist[page])
        # else: tome - 1
        tomelist = (path.strip(self.book) for path in os.listdir(f'xml') if path.startswith(self.book))
        tomelist = [f'{self.book}{i}' for i in sorted(tomelist, key=int)]
        idx = tomelist.index(f'{self.book}{self.tome}')
        if 0 <= idx - 1:
            nextbook = tomelist[idx - 1]
            juanlist = get_sorted_juan(nextbook)
            return Number(juanlist[page])
        # else: book - 1
        return juanlist

    @property
    def title(self):
        number = f'{self.book}{self.tome}n{self.sutra}{self.yiyi}'
        title = ''
        with open("idx/sutra_sch.lst") as fd:
            for line in fd:
                if number in line:
                    title = line.strip().split(maxsplit=1)[1]
                    break
        return title


def get_sorted_juan(book):
    '''获得book(T01)下的所有排序好的经卷号(T01n0002_001)'''
    # 对所有的book下的卷排序
    juanlist = []
    for path in os.listdir(f'xml/{book}'):
        sutra, juan = path[:-4].split('_')
        sutra = sutra.split('n')[1]
        juanlist.append((sutra, juan))
    juanlist.sort(key=lambda x: (int(f'{x[0]:0<5}', 16), int(x[1])))
    juanlist = [f'{book}n{i[0]}_{i[1]:0>3}' for i in juanlist]
    return juanlist


def grep(filepath, *keyword):
    '''对Linux命令grep的模拟，返回一行'''
    found = False
    with open(filepath) as fd:
        for line in fd:
            line = line.strip()
            if all(kw in line for kw in keyword):
                found = True
                break
    if not found:
        return None
    return line

sch_db = []
#with open(os.path.join(PATH, "idx/sutra_sch.lst")) as fd:
with open("idx/sutra_sch.lst") as fd:
    for line in fd:
        line = line.strip().split()[0]
        sch_db.append(line)

juan_pattern = re.compile(r'(\s+)第?([\d零〇一二三四五六七八九]{1,4})卷?')

# ../T31/T31n1585.xml#xpath2(//0011b12)
# 雜阿含經一五·一七
# 增一阿含二一·六（大正二·六〇三c）  <pb n="0603c" ed="T" xml:id="T02.0125.0603c"/>
ahan_pattern = re.compile(r'(《?[中长長雜杂增][一壹]?阿[含鋡][經经]?》?)第?([\d零〇一二三四五六七八九]{1,4})[經经]')
def parse_ahan(number):
    found = False
    jinghao = ahan_pattern.findall(number)
    if jinghao:
        book, sutra = jinghao[0]
        sutran = int(sutra)
        if '雜' in book or '杂' in book:
            book = 'T0099'
            if sutran > 1362:
                sutran = 1362
            if sutran < 1:
                sutran = 1
        else:
            return None

        # 查表
        with open(f'idx/{book}.xml') as fd:
            for line in fd:
                line = line.strip().split()
                if sutran == int(line[0]):
                    sutra = Number(line[1])
                    found = True
                    break
    if not found:
        return None
    return sutra


# 大正七〇·四五九中、四六〇下
# 大正二、一〇三c
# 大正藏二·四一a
# 大正藏第70卷459页b
# 《大正藏》第40卷第16頁下
# 大正藏第十九卷第16頁下 XXX
pbanchor_pattern = re.compile(r'(《?[中乾佛作傳典刊刻北卍南印叢史品善嘉國圖城外大學宋家寺山師彙志房拓教文新書朝本樂正武永法洪漢片獻珍百石經編纂續脩興華著藏補譯趙遺金隆集順館高麗传丛国图学师汇书乐汉献经编续修兴华补译赵遗顺馆丽]+》?)第?([\d零〇一二三四五六七八九]{1,3})(?:卷|卷第|\u00b7)([\d零〇一二三四五六七八九]{1,3})[頁|页]?([上中下abcABC])?')
def parse_number2(number):
    tt = { ord('〇'): ord('0'), ord('零'): ord('0'),
          ord('一'): ord('1'),
          ord('二'): ord('2'),
          ord('三'): ord('3'),
          ord('四'): ord('4'),
          ord('五'): ord('5'),
          ord('六'): ord('6'),
          ord('七'): ord('7'),
          ord('八'): ord('8'),
          ord('九'): ord('9'),
          ord('上'): ord('a'),
          ord('中'): ord('b'),
          ord('下'): ord('c'),
          ord('伍'): ord('5'),
          ord('叁'): ord('3'), ord('叄'): ord('3'),
          ord('壹'): ord('1'),
          ord('捌'): ord('8'),
          ord('柒'): ord('7'),
          ord('玖'): ord('9'),
          ord('肆'): ord('4'),
          ord('貳'): ord('2'), ord('贰'): ord('2'),
          ord('陆'): ord('6'), ord('陸'): ord('6'),
          # 伍叁叄壹拾捌柒玖肆貳贰陆陸
            }
    number = re.sub(r'\s+', '', number)
    found = False
    book = 'T'
    anchor = ''
    jinghao = pbanchor_pattern.findall(number)
    print(number, jinghao)
    if jinghao:
        book, tome, page, abc = jinghao[0]
        page = '{:04}'.format(int(page.translate(tt)))
        abc = abc.translate(tt).lower()
        if '大正' in book:
            book = 'T'
        if '卍新' in book or '卍續' in book or '卍续' in book:
            book = 'X'
        if '高麗' in book or '高丽' in book:
            book = 'K'
        if '房山' in book:
            book = 'F'
        if '印順' in book or '印顺' in book:
            book = 'Y'
        if '宋藏' in book:
            book = 'S'
        if '金' in book:
            book = 'A'
        if '中華' in book or '中华' in book:
            book = 'C'
        if '嘉興' in book or '嘉兴' in book:
            book = 'J'
        if '永樂北' in book or '永乐北' in book:
            book = 'P'
        if '洪武南' in book:
            book = 'U'
        if '圖' in book:
            book = 'D'
        if '南傳' in book or '南传' in book:
            book = 'N'
        if '藏外' in book:
            book = 'ZW'
        if '補' in book or '补' in book:
            book = 'B'
        if '乾隆' in book:
            book = 'L'
        if '呂澂' in book or '吕澂' in book:
            book = 'LC'
        if book in {'A', 'C', 'G', 'GA', 'GB', 'L', 'M', 'P', 'U'}:
            tome = '{:03}'.format(int(tome.translate(tt)))
        else:
            tome = '{:02}'.format(int(tome.translate(tt)))

        # 查表
        with open('idx/pbidx.txt') as fd:
            for line in fd:
                line = line.strip()
                if f'{book}{tome}' in line and f'{page}{abc}' in line:
                    number, anchor = line.split()
                    found = True
    if not found:
        return None
    # url = f'/xml/{book}{tome}/{number}.xml#{anchor}'
    # return url
    return Number(f'{number}#{anchor}')
    # return (book, tome, sutra, j4, volume, anchor)

# 标准模式: T01n0001, T01n0001_001
# 模式1: T01n0001_p0001a01
# 模式2: T01,no.1,p.1a1
# CBETA 2019.Q2, Y25, no. 25, p. 411a5-7
# CBETA, T14, no. 475, pp. 537c8-538a14
# CBETA 2019.Q3, T20, no. 1113B, p. 498c12-17
# CBETA 2019.Q3, A091, no. 1057, pp. 311b03-312a10
# 模式9: '100,3', 't100,3', 100, t1000, t1000_001, 1333b
# TODO: 大宝积经100
jinghaopatten = re.compile(r'([A-Z]{1,2})(\d{2,3})n(\w\d{3})([a-zA-Z])?(?:_(\d{3}))?')
jinghaopatten1 = re.compile(r'([a-zA-Z]{1,2})(\d{2,3})n(\w\d{3})([a-zA-Z])?(?:_(\d{3}))?(?:_p(\w\d{3}[abc]\d\d))?')
jinghaopatten2 = re.compile(r'([a-zA-Z]{1,2})(\d{2,3}),\s*no\.\s*(\w\d{0,3})([a-zA-Z])?,\s*pp?\.\s*(\d+)([abc])(\d+)')
jinghaopatten9 = re.compile(r'([a-zA-Z]{1,2})?(\w\d{0,3})([a-zA-Z])?(?:[\s,._\u3002\uff0c-]+(\d+)?)?')  # 全角逗号句号
# jinghaopatten3 = re.compile(r'([\u3007\u3400-\u9FCB\U00020000-\U0002EBE0]+)[ \t,._\u3000\u3002\uff0c-]*(\d+)')
def parse_number1(title, guess_juan=False):
    '''所有的anchor都是lb标签的'''
    # print(title)
    book, tome, sutra, j4, volume, anchor = 'T', '', '', '', '', ''
    # book, tome, sutra, j4, volume, anchor
    #    T,   01,  0001,  a     001, p0001a01
    found = False
    if not found:
        jinghao = jinghaopatten1.findall(title)
        if jinghao:
            book,tome,sutra,j4,volume,anchor = jinghao[0]
            found = True

    if not found:
        jinghao = jinghaopatten2.findall(title)
        if jinghao:
            book,tome,sutra,j4,anchor,j7,j8 = jinghao[0]
            anchor = '{:04}{}{:02}'.format(int(anchor), j7, int(j8))
            found = True

    if not found:
        jinghao = jinghaopatten9.findall(title)
        if jinghao:
            book,sutra,j4,volume = jinghao[0]
            found = True

    if title.isdigit():
        sutra = '{:04}'.format(int(title))
        found = True

    if not found:
            return None

    book = book.upper() if book else 'T'
    sutra = sutra.upper().zfill(4)
    # TODO: 查找卷数
    if not tome:
        # 大般若经特例
        if volume and f'{book}{sutra}' == 'T0220':
            if int(volume) <= 200:
                tome = '05'
            if 201 <= int(volume) <= 400:
                tome = '06'
            if int(volume) > 400:
                tome = '07'
        else:
            for line in sch_db:
                if book in line and f'n{sutra}' in line:
                    tome = line.split('n')[0][len(book):]
                    break
    # print(1, '|'.join((title, book,tome,sutra,j4,volume,anchor)))
    if not tome:
        return None

    # 用book,tome,sutra确定j4;有j4确定大小写
    found = False
    # j4如果是小写就变为大写, 大写就变成小写
    j9 = j4.upper() if j4 and ord('a') <= ord(j4) <= ord('z') else j4.lower()
    for line in sch_db:
        if f'{book}{tome}n{sutra}{j4}' in line:
            j4 = line[len(f'{book}{tome}n{sutra}'):]
            found = True
            break
        if f'{book}{tome}n{sutra}{j9}' in line:
            j4 = line[len(f'{book}{tome}n{sutra}'):]
            found = True
            break

    if not found:
        return None

    # 查找第一卷的卷数
    if guess_juan and not volume and not anchor:
        volume = get_first_juan(f'{book}{tome}n{sutra}{j4}')

    # 根据锚来查找卷数volume
    if not volume and anchor:
        with gzip.open(f'idx/lb/{book}{tome}n{sutra}{j4}.gz', 'rt') as fd:
            for line in fd:
                if anchor in line:
                    volume = line.split()[0].split('_')[-1]
                    break

        # print(4, title, book, tome,sutra,j4,volume, anchor)

    # 没找到,非法字符串
    if guess_juan and not volume:
        return None

    if volume:
        volume = int(volume)
    else:
        volume = 0
    #    volume = '{:03}'.format(int(volume))
    # ('T', '05', '0220', '', 0, '')
    return Number((book, tome, sutra, j4, volume, anchor))


def parse_number(number):
    '''尝试解析不同的经号格式'''
    sutra = parse_number1(number)
    if sutra:
        return sutra
    # 使用卷册方式查找藏经
    sutra = parse_number2(number)
    if sutra:
        return sutra
    # 查找阿含经
    sutra = parse_ahan(number)
    return sutra

# def normalize_number(number, guess_juan=False):
#     '''如果number符合经号的形式, 就标准化为标准形式T01n0001_001, 否则原样返回'''
#     result = parse_number1(number, guess_juan)
#     if result:
#         j1, j2, j3, j4, j5, j6 = result
#         if j5:
#             number = f'{j1}{j2}n{j3}{j4}_{j5}'
#         else:
#             number = f'{j1}{j2}n{j3}{j4}'
#     return number
#
#
# def make_url(number):
#     number = parse_number1(number, True)
#     # 如果有锚就添加锚
#     if number:
#         j1, j2, j3, j4, j5, j6 = number
#     else:
#         return None
#     if j6:
#         url = f'xml/{j1}{j2}/{j1}{j2}n{j3}{j4}_{j5}.xml#{j6}'
#     else:
#         url = f'xml/{j1}{j2}/{j1}{j2}n{j3}{j4}_{j5}.xml'
#     return url

# FROM: https://en.wikipedia.org/wiki/International_Alphabet_of_Sanskrit_Transliteration

# FROM: https://en.wikipedia.org/wiki/Harvard-Kyoto
class SA:
    '''梵语字符串类, 可以使用HK转写和iast转写输入, 使用天城体, 悉檀体, 拉丁输出'''
    pass

def fromlatn(ctx):
    vowel = {
            'a': chr(0x11580), 'ā': chr(0x11581), 'i': chr(0x11582), 'ī': chr(0x11583),
            'u': chr(0x11584), 'ū': chr(0x11585), 'ṛ': chr(0x11586), 'ṝ': chr(0x11587),
            'ḷ': chr(0x11588), 'ḹ': chr(0x11589), 'e': chr(0x1158a), 'ai': chr(0x1158b),
            'o': chr(0x1158c), 'au': chr(0x1158d),
            }
    consonant = {
    'k': chr(0x1158e), 'kh': chr(0x1158f), 'g': chr(0x11590), 'gh': chr(0x11591),
    'ṅ': chr(0x11592), 'c': chr(0x11593), 'ch': chr(0x11594), 'j': chr(0x11595),
    'jh': chr(0x11596), 'ñ': chr(0x11597), 'ṭ': chr(0x11598), 'ṭh': chr(0x11599),
    'ḍ': chr(0x1159a), 'ḍh': chr(0x1159b), 'ṇ': chr(0x1159c), 't': chr(0x1159d),
    'th': chr(0x1159e), 'd': chr(0x1159f), 'dh': chr(0x115a0), 'n': chr(0x115a1),
    'p': chr(0x115a2), 'ph': chr(0x115a3), 'b': chr(0x115a4), 'bh': chr(0x115a5),
    'm': chr(0x115a6), 'y': chr(0x115a7), 'r': chr(0x115a8), 'l': chr(0x115a9),
    'v': chr(0x115aa), 'ś': chr(0x115ab), 'ṣ': chr(0x115ac), 's': chr(0x115ad),
    'h': chr(0x115ae),
    }
    mvowel = {
    'ā': chr(0x115af),
    'i': chr(0x115b0),
    'ī': chr(0x115b1),
    'u': chr(0x115b2),
    'ū': chr(0x115b3),
    'r': chr(0x115b4),
    'ṛ': chr(0x115b5),
    'e': chr(0x115b8),
    'ai': chr(0x115b9),
    'o': chr(0x115ba),
    'au': chr(0x115bb),
    }
    sign = {
    'ṃ': chr(0x115bd),
    'ḥ': chr(0x115be),
    }
    # '|': chr(0x115c2),
    # '||': chr(0x115c3),


def hk2iastdeve(str_in):
    '''hk哈佛-京都系统转IAST梵语(天城体)'''
    sonorants  = {
            # Sonorants:
            'RR': 'ॠ',
            'lR': 'ऌ',
            'lRR': 'ॡ ',
        }

    t1 = {
            # Anusvāra and visarg:
            'aM': 'अं',
            'aH': 'अः',
        # Consonants:
            'ai': 'ऐ',
            'au': 'औ',
            'kh': 'ख',
            'gh': 'घ',
            'ch': 'छ',
            'jh': 'झ',
            'ph': 'फ',
            'Th': 'ठ',
            'Dh': 'ढ',
            'th': 'थ',
            'dh': 'ध',
            'bh': 'भ',
            }

    t2 = {
            'R': 'ऋ',
            # Vowels:
            'a': 'अ',
            'A': 'आ',    # ā
            'i': 'इ',
            'I':'ई',
            'u': 'उ',
            'U':'ऊ',
            'e': 'ए',
            'o': 'ओ',

        # Consonants:
            'k': 'क',
            'g': 'ग',
            'G': 'ङ',
            'c': 'च',
            'j': 'ज',
            'J': 'ञ',
            'T': 'ट',
            'D': 'ड',
            'N': 'ण',
            't': 'त',
            'd': 'द',
            'n': 'न',
            'p': 'प',
            'b': 'ब',
            'm': 'म',
            'y': 'य',
            'r': 'र',
            'l': 'ल',
            'v': 'व',
            'z': 'श',
            'S': 'ष',
            's': 'स',
            'h': 'ह',
            }

        # '@':' ',

    # usedt = {ord(k): ord(t1[k]) for k in t1}
    str_out = str_in.replace('RR', sonorants['RR']).replace('lR', sonorants['lR']).replace('lRR', sonorants['lRR'])
    for zi in t1:
        str_out = str_out.replace(zi, t1[zi])
    for zi in t2:
        str_out = str_out.replace(zi, t2[zi])
    # str_out = str_out.translate(usedt)
    return str_out


def hk2iast(str_in):
    '''hk哈佛-京都系统转IAST梵语'''
    x = {'S':'sh',
        'RR':'\u1e5b\u012b'}  #   1e5d
        # 'RR':'\u1e5d'}     # 1e5d
        # 'lR':'\u1eca'}     # 1e5d
        # 'lRR':'\u1e39'}     # 1e5d

    t1 = {'A': '\u0101',    # ā
        'I':'\u012b',
        'U':'\u016b',
        'M':'\u1e43', # 1e43
        'H':'\u1e25',
        'G':'\u1e45',
        'J':'\u00f1',
        'T':'\u1e6d',
        'D':'\u1e0d',
        'N':'\u1e47',
        'L':'\u1eca',   # Ị
        'z':'\u1e61',
        '@':' ',
        'R':'\u1e5b',      # ṛ
        'S':'\u1e63',      #
        }

    t2 = {'A': '\u0101',
        'I':'\u012b',
        'U':'\u016b',
        'M':'\u1e49', # 1e49
        'H':'\u1e25',
        'G':'\u1e45',
        'J':'\u00f1',
        'T':'\u1e6d',
        'D':'\u1e0d',
        'N':'\u1e47',
        'L':'\u1eca',
        'z':'\u1e61',
        '@':' ',
        }

    usedt = {ord(k): ord(t1[k]) for k in t1}
    str_out = str_in.replace('RR', '\u1e5d').replace('lR', '\u1eca').replace('lRR', '\u1e39')
    str_out = str_out.translate(usedt)
    return str_out

hk2sa = hk2iast

def HKdict2iast(hkdict):
    # 将HK系统梵文词典转为IAST系统梵文词典
    mwpatten = re.compile(r'(%\{.+?})')
    sa_en = dict()
    # for key in data:
    #     k = key.replace('1', '').replace("'", '').replace('4', '').replace('7', '').replace('8', '').replace('9', '').replace('0', '').replace('-', '').lower()
    #     sa_en.update({k: data[key]})

    for key in hkdict:
        vals = hkdict[key]
        devkey = hk2iastdeve(key)
        key = hk2iast(key)  # .replace('1', '').replace("'", '').replace('4', '').replace('7', '').replace('8', '').replace('9', '').replace('0', '').replace('-', '') #.lower()
        # 将天城体附加在罗马体后面
        key = ' '.join((key, devkey))
        res = []
        for val in vals:
            x = mwpatten.findall(val)
            if x:
                for ff in x:
                    val = val.replace(ff, hk2iast(ff))
            res.append(val)
        # 不知道以下这两行那个对
        sa_en.update({key: res})
    return sa_en


# def load_dict(dictionary=None):
#
#     # 词典列表
#     dicts = {'fk': ('佛光山', 'fk.json.gz'), 'dfb': ('丁福保', 'dfb.json.gz'), 'ccc': ('庄春江', 'ccc.json'), 'fxcd': ('法相詞典', 'fxcd.json.gz'),
#             'nvd': ('南山律学词典', 'nvd.json.gz'), 'cyx': ('佛學常見詞彙（陳義孝）', 'cyx.json'), 'ylb': ('唯识名词白话新解', 'ylb.json'),
#             'szfs': ('三藏法数', 'szfs.json'), 'fymyj': ('翻譯名義集', 'fymyj.json'), 'wdhy': ('五燈會元', 'wdhy.json.gz'), 'yzzj': ('閱藏知津', 'yzzj.json.gz'),
#             'ldms': ('歷代名僧辭典', 'ldms.json.gz'), 'syfy': ('俗語佛源', 'syfy.json.gz'), 'bkqs': ('中华佛教百科全书','bkqs.json.gz')}
#
#     aio = dict()
#
#     # 装入梵英词典, 太大了，暂时不装了
#     mwpatten = re.compile(r'(%\{.+?})')
#     sa_en = dict()
#
#     # s = time.time()
#     # with gzip.open('dict/sa-en.json.gz') as fd:
#     #     data = fd.read()
#     # data = json.loads(data)
#     # sa_en = dict()
#     # for key in data:
#     #     k = key.replace('1', '').replace("'", '').replace('4', '').replace('7', '').replace('8', '').replace('9', '').replace('0', '').replace('-', '').lower()
#     #     sa_en.update({k: data[key]})
#     #
#     # for key in data:
#     #     vals = data[key]
#     #     res = []
#     #     for val in vals:
#     #         x = mwpatten.findall(val)
#     #         if x:
#     #             for ff in x:
#     #                 val = val.replace(ff, hk2sa(ff))
#     #         res.append(val)
#     #     # 不知道以下这两行那个对
#     #     sa_en.update({hk2sa(key, 1): res})
#     #     sa_en.update({hk2sa(key, 2): res})
#     # e = time.time()
#     # print('装入梵英词典，用时%s' % (e - s))
#     yield ('sa_en', sa_en)
#
#     sa_hant = dict()
#     # s = time.time()
#     # with gzip.open('dict/sa-hant.json.gz') as fd:
#     #     data = fd.read()
#     # data = json.loads(data)
#     # for key in data:
#     #     sa_hant.update({key.lower(): data[key]})
#     # e = time.time()
#     # print('装入梵汉词典，用时%s' % (e - s))
#     yield ('sa_hant', sa_hant)
#
#     yat = dict()
#     # s = time.time()
#     # with gzip.open('dict/yat.json.gz') as fd:
#     #     data = fd.read()
#     # data = json.loads(data)
#     # for key in data:
#     #     yat.update({key.lower(): data[key]})
#     # for key in data:
#     #     vals = data[key]
#     #     res = []
#     #     for val in vals:
#     #         x = mwpatten.findall(val)
#     #         if x:
#     #             for ff in x:
#     #                 v = val.replace(ff, hk2sa(ff))
#     #         res.append(v)
#     #     yat.update({hk2sa(key, 1): res})
#     #     yat.update({hk2sa(key, 2): res})
#     # e = time.time()
#     # print('装入Yates梵英词典，用时%s' % (e - s))
#     yield ('yat', yat)
#
#     s = time.time()
#     with gzip.open('dict/kangxi.json.gz') as fd:
#         kangxi = json.load(fd)
#     e = time.time()
#     print('装入康熙字典，用时%s' % (e - s))
#     yield ('kangxi', kangxi)
#
#     s = time.time()
#     with open('dict/Unihan_Readings.json') as fd:
#         unihan = json.load(fd)
#     e = time.time()
#     print('装入Unicode10.0字典，用时%s' % (e - s))
#     yield ('unihan', unihan)
#
#     for k in dicts:
#         s = time.time()
#         path = f'dict/{dicts[k][1]}'
#         if not os.path.exists(path):
#             continue
#
#         if path.endswith('gz'):
#             with gzip.open(path) as fd:
#                 try:
#                     v = json.load(fd)
#                 except:
#                     print(path)
#                     raise
#         else:
#             with open(path, encoding='utf8') as fd:
#                 v = json.load(fd)
#         # for k1 in v:
#         #     if k1 in aio:
#         #         aio[k1].update({dicts[k][0]: v[k1]})
#         #     #     aio[k1].append()
#         #     else:
#         #         aio[k1] = {dicts[k][0]: v[k1]}
#         e = time.time()
#         print('装入%s，用时%s' % (dicts[k][0], e - s))
#         yield (k, v)
#
#     yield ('aio', aio)
#
#     # return {'kangxi':kangxi, 'unihan':unihan,
#     #         'fk':fk, 'dfb': dfb, 'ccc': ccc, 'nvd': nvd, 'cxy': cxy, 'ylb': ylb, 'fxcd': fxcd,
#     #         'szfs': szfs, 'fymyj': fymyj,
#     #     'sa_hant': sa_hant, 'sa_en': sa_en, 'yat': yat}
#
#
#
#
# def lookinsa(word):
#     definition = sa_hant.get(hk2sa(word).lower(), '')
#     pinyin = ""
#     _from = ""
#     if definition:
#         _from = "文理学院"
#         pinyin = "文理学院"
#     if not definition:
#         # 使用Harvard-Kyoto转写查找字典
#         definition = sa_en.get(hk2sa(word), '')
#         # 使用缩写查找字典
#         if not definition:
#             w = word.replace('1', '').replace("'", '').replace('4', '').replace('7', '').replace('8', '').replace('9', '').replace('0', '').replace('-', '').lower()
#             definition = sa_en.get(w, '')
#         if definition:
#             definition = '|'.join(definition)
#             _from = "威廉梵英词典"
#             pinyin = "威廉梵英词典"
#     if not definition:
#         print(hk2sa(word))
#         definition = yat.get(hk2sa(word), '')
#         if not definition:
#             w = word.replace('-', '').lower()
#             definition = yat.get(w, '')
#         if definition:
#             definition = '|'.join(definition)
#             _from = "YAT"
#             pinyin = "YAT"
#     return {'word': word, 'pinyin': pinyin, 'def': definition, 'from': _from}


# 装入各种词典
# dd = dict(load_dict())
# sa_hant = dd['sa_hant']
# sa_en = dd['sa_en']
# yat = dd['yat']
# kangxi = dd['kangxi']
# unihan = dd['unihan']
# fk = dd['fk']
# dfb = dd['dfb']
# ccc = dd['ccc']
# nvd = dd['nvd']
# cyx = dd['cyx']
# ylb = dd['ylb']
# fxcd = dd['fxcd']
# szfs = dd['szfs']
# fymyj = dd['fymyj']
# wdhy = dd['wdhy']
# ldms = dd['ldms']
# yzzj = dd['yzzj']
# bkqs = dd['bkqs']

# aio = dd['aio']


class Search:
    def __init__(self, norm=True):
        titles = []
        with open("idx/sutra_sch.lst") as fd:
            for line in fd:
                line = normalize_text(line.strip()).split(maxsplit=1)
                titles.append(line)

        # pprint.pprint(titles)
        # titles 是经号和title的对照表
        # 生成索引表
        self.index = {}
        for i in titles:
            z = 0
            #print(i)
            for j in i[1]:
                if j in self.index:
                    self.index[j].append((i[0], z))
                else:
                    self.index[j] = [(i[0], z),]
                #print(j, i[0], z)
                z += 1
        for i in self.index:
            # print(i)
            v = self.index[i]
            r = dict()
            for j in v:
                if j[0] in r:
                    r[j[0]].append(j[1])
                else:
                    r[j[0]] = [j[1],]
            # pprint.pprint((i, r))
            self.index.update({i: r})
        self.titles = dict(titles)

    def search(self, title, norm=True):
        '''搜索经文标题或者作者，以标题为主'''
        # title = opencc.convert(title, config='s2t.json')
        # ( for zi in index)
        if norm:
            title = normalize_text(title)
        result = (set(self.index.get(tt, {}).keys()) for tt in list(title))
        # return sorted(reduce(lambda x, y: x & y, result), key=pagerank)
        result = reduce(lambda x, y: x & y, result)
        #return sorted(result, key=lambda x: Levenshtein.ratio(title, self.titles[x].split(' (')[0]), reverse=True)
        return sorted(result, key=lambda x: (1 if x[0] == 'T' else 0, Levenshtein.ratio(title, self.titles[x].split(' (')[0])), reverse=True)


# 简体繁体转换
def re_split(pattern, string, fn=lambda x: x, exfn=lambda x: x):
    '''对re模块split的改进；把输入字符串使用pattern分割, 匹配的字符串使用fn处理, 不匹配的使用exfn处理'''
    if not isinstance(pattern, re.Pattern):
        pattern = re.compile(pattern)
    rr = pattern.search(string)
    while rr:
        start, end = rr.span()
        if start != 0:
            yield exfn(string[:start])
        yield fn(string[start:end])
        string = string[end:]
        rr = pattern.search(string)
    if string:
        yield exfn(string)


class STConvertor:
    '''简体繁体转换类'''
    def __init__(self):
        '''读取简体繁体转换数据库'''
        # 读取繁体转简体短语词典
        self.tsptable = readdb('cc/TSPhrases.txt')
        # 读取简体转繁体短语词典
        self.stptable = readdb('cc/STPhrases.txt')
        # # print('|'.join(sorted(tsptable.keys(), key=lambda x: len(x), reverse=True)))
        # 读取繁体转简体字典
        self.tstable = readdb('cc/TSCharacters.txt', trans=True)
        # 读取简体转繁体字典
        self.sttable = readdb('cc/STCharacters.txt', trans=True)

        # 简体繁体转换pattern
        self.tsp = re.compile('|'.join(sorted(self.tsptable.keys(),key=len,reverse=True)))
        self.stp = re.compile('|'.join(sorted(self.stptable.keys(),key=len,reverse=True)))

        # return tsp, tstable, tsptable, stp, sttable, stptable

        # 简体繁体检测
        # self.p = re.compile(r'[\u4e00-\u9fa5]')

        _tst = readdb('cc/TSCharacters.txt')
        # self.tt: 纯繁体字集合
        # self.ss: 纯简体字集合
        tt = set(_tst.keys())
        ss = set(_tst.values())
        xx = tt & ss

        self.tt = tt - xx
        self.ss = ss - xx

        self.jt = set()
        with open('idx/jt.txt') as fd:
            for line in fd:
                if line.startswith('#'): continue
                line = line.split()
                self.jt.add(line[0])

    def t2s(self, string, punctuation=True, region=False, autonorm=True, onlyURO=True):
        '''繁体转简体, punctuation是否转换单双引号
        region 是否执行区域转换
        region 转换后的地区
        autonorm 自动规范化异体字
        onlyURO 不简化低位类推简化字(繁体字处于BMP和扩展A区, 但是简体字处于扩展B,C,D,E,F的汉字)
        '''
        if autonorm:
            string = rm_variant(string)

        if punctuation:
            string = string.translate({0x300c: 0x201c, 0x300d: 0x201d, 0x300e: 0x2018, 0x300f: 0x2019})

        # 类推简化字处理
        tst2 = copy.deepcopy(self.tstable)
        if onlyURO:
            # 只要简化字不在BMP，就是类推简化字
            # tst2 = {k:tstable[k] for k in tstable if not (k < 0x20000 and tstable[k] > 0x20000)}
            # tst2 = {k:tstable[k] for k in tstable if 0x4E00 <= tstable[k] < 0x20000}
            tst2 = {k:self.tstable[k] for k in self.tstable if 0x4E00 <= self.tstable[k] < 0x9FA5}
        else:
            tst2 = {k:self.tstable[k] for k in self.tstable}

        content = ''.join(re_split(self.tsp, string, self.tsptable.get, lambda i: i.translate(tst2)))

        return content


    def s2t(self, string, punctuation=True, region=False):
        '''简体转繁体, punctuation是否转换单双引号
        region 是否执行区域转换
        region 转换后的地区
        '''

        if punctuation:
            string = string.translate({0x201c: 0x300c, 0x201d: 0x300d, 0x2018: 0x300e, 0x2019: 0x300f})

        content = ''.join(re_split(self.stp, string, self.stptable.get, lambda i: i.translate(self.sttable)))

        return content


    # def detect(self, s0):
    #     '''粗略判断一段文本是简体还是繁体的概率'''
    #     if len(s0) == 0:
    #         return {'t': 50, 's': 50, 'confidence': 's'}

    #     s0 = set(s0)
    #     # 同时是简体繁体的可能性
    #     j = sum(1 for i in (s0 - self.tt - self.ss) if self.p.match(i))
    #     # 繁体可能性
    #     t = 100 + ((j * 50 - len(s0 - self.tt) * 100 )/ len(s0))
    #     # 简体可能性
    #     s = 100 + ((j * 50 - len(s0 - self.ss) * 100 )/ len(s0))

    #     confidence = 's'
    #     if t > 50:
    #         confidence = 't'
    #     elif s > 50:
    #         confidence = 's'
    #     elif t > s:
    #         confidence = 't'

    #     return {'t': t, 's': s, 'confidence': confidence}

    def detect(self, s0):
        '''使用简体字表来判断一段文本是简体还是繁体的概率'''
        if len(s0) == 0:
            return {'t': 50, 's': 50, 'confidence': 's'}
        t = 50
        s = 50
        confidence = 't'
        for zi in s0:
            if zi in self.jt:
                confidence = 's'
                break

        return {'t': t, 's': s, 'confidence': confidence}


# 异体字处理

# 读入异体字对照表
yitizi = readdb('dict/variants.txt', True, True)
# 读取异体字短语词典
varptable = readdb('variants/p.txt')
# 异体短语转换pattern
varppp = re.compile('|'.join(sorted(varptable.keys(),key=len,reverse=True)))

def rm_variant(string, level=0):
    '''异体字规范化为标准繁体字'''
    # ctx = unicodedata.normalize("NFKC", ctx)
    content = ''.join(re_split(varppp, string, varptable.get, lambda i: i.translate(yitizi)))
    return content


def zi_order(ss, ct):
    '''判断ss字符串中的字是否按照顺序在ct字符串中出现'''
    rr = dict()
    for i, zi in enumerate(ct):
        if zi in rr:
            rr[zi].add(i)
        else:
            rr[zi] = {i}

    result = []
    for zi in ss:
        if zi not in rr:
            return False
        else:
            result.append(rr[zi])

    start = -1
    for i in result:
        start = {j for j in i if j > start}
        if not start:
            return False
        else:
            start = min(start)
    return True


def highlight(ss, ct):
    for zi in set(ss):
        ct = ct.replace(zi, f'<em>{zi}</em>')
    return ct


# pun = string.punctuation + '\u00b7\u2013-\u2027\u2e3a\u3000-\u301f\ufe30-\ufe6b\uff01-\uff0f\uff1a-\uff5e'
# pun = re.compile('['+string.punctuation+']')
# 读取标点数据库
pun = dict()
with open('dict/punctuation.txt') as fd:
    for line in fd:
        line = line.strip()
        if line.startswith('~~~~~~~~~'):
            break
        if not line or line.startswith('#'): continue
        c1 = line[0]
        pun[ord(c1)] = 0xFFFD


def rm_pun(ctx, ex=()):
    '''删除标点符号，除了ex列表中的字符'''
    for char in ex:
        pun.pop(ord(char), None)
    ctx = ctx.translate(pun).replace(chr(0xFFFD), '')
    return ctx


def pagerank(filename, sentence='', content=''):
    '''对xml文件名评分, filename 为 T20n1060 或者 T20n1060_001.xml 形式
    A,B,C,D,F,G,GA,GB,I,J,K,L,M,N,P,S,T,U,X,ZW, Y, LC
    '''
    # sentence = sentence.strip().split()
    # sentence_value = sum([{True:0, False:1}[s in content] for s in sentence])
    pr = ("T", "B", "ZW", "A", "C", "D", "F", "G" , "GA", "GB", "I", "J", "K", "L", "M", "N", "P", "S", "U", "X", "Y", "LC")
    pt = re.compile(r'\d+')  # 应该在前端过滤
    if filename[0] == 'T':
        r = 0
    else:
        r = 1
    x = pt.findall(filename)
    x = [int(i) for i in x]
    # return r, x[0], x[1], x[2]
    x.insert(0, r)
    # x.insert(0, sentence_value)
    # 对于符合条件的重新排列到前面
    # zi_order(ss, ct)
    return x



def must_search(sentence, _from=0, _end=5000):
    print('must搜索:', sentence)
    url = "http://127.0.0.1:9200/cbeta/_doc/_search" #创建一个文档，如果该文件已经存在，则返回失败
    data = {
     "query": {
        # "match_phrase": { "content": sentence # "content": {"query": sentence, "slop": 1} },
         "bool":{
            #  "must": {}
        }
    },
    "size":_end - _from ,
    "from": _from,
    # "highlight": {
    #     "fields": {
    #         "raw": {

    #         }
    #     }
    # }
    }


    if re.findall(r'\s+and\s+|\s*&\s*', sentence, flags=re.I):
        sentences = re.split(r'\s+and\s+|\s*&\s*', sentence, flags=re.I)
        sentences = [re.sub(r'\s+', '', ctx) for ctx in sentences]
    else:
        sentences = re.split(r'\s+', sentence, flags=re.I)
    # data["query"]["bool"]["must"] = [{"match_phrase": {"content": st}} for st in sentences]
    sentences = [re.split(r':|：', st) for st in sentences]
    # 组成标准查询JSON
    # must = [("content", st[0]) if len(st) == 1 else (st[0].lower(), st[1] if st[0] != 'number' else normalize_number(st[1],False)) for st in sentences]
    # must = [('sutra' if st0=='number' and '_' not in st1 else st0, st1) for st0, st1 in must]
    # data["query"]["bool"]["must"] = [{"match_phrase": {"content": val}} if key=="content" else {"match": {key: val}} for key,val in must]
    data["query"]["bool"]["must"] = [{"match_phrase": {"content": st[0]}} if len(st) == 1 else {"match": {st[0].lower(): st[1]}} for st in sentences]
    pprint.pprint(data["query"]["bool"]["must"])

    r = requests.get(url, json=data, timeout=10)
    result = r.json()
    return result


def wordsearch(word):
    # nword = normalize_text(word)
    # sentence = ''.join(python_escape(sentence))
    url = "http://127.0.0.1:9200/dict/_doc/_search"
    data = {
     "query": {
         # "match_phrase": { "orth": nword},  # "content": {"query": sentence, "slop": 1} },
         "match": {"orth": word},  # "content": {"query": sentence, "slop": 1} },
       #  "bool":{
            #  "must": {}
       # }
    },
    "size": 5000,
    "from": 0,
    # "highlight": {
    #     "fields": {
    #         "raw": {

    #         }
    #     }
    # }
    }
    r= requests.get(url, json=data, timeout=10)
    result = r.json()
    # print(result)
    hits = result['hits']['hits']
    #print(hits)
    result = []
    for hit in hits:
        _source = hit["_source"]
        # 文章内容高亮显示
        result.append({'hl': _source['def'], 'an': '', 'title':_source['hyph'], 'author': _source['dict'], 'number': ''})

    # result.sort(key=lambda x: pagerank(x['number']))
    return result


def fullsearch(sentence):
    '''全文搜索, sentence是繁体字'''
    # 标准化文本
    sentence = normalize_text(sentence)
    sentence = ''.join(python_escape(sentence))
    url = "http://127.0.0.1:9200/cbeta/_doc/_search"
    data = {
     "query": {
        # "match_phrase": { "content": sentence # "content": {"query": sentence, "slop": 1} },
         "bool":{
            #  "must": {}
        }
    },
    "size": 5000,
    "from": 0,
    # "highlight": {
    #     "fields": {
    #         "raw": {

    #         }
    #     }
    # }
    }

    # 组成标准查询JSON
    s = time.time()
    if re.findall(r'\s+and\s+|\s*&\s*', sentence, flags=re.I):
        sentences = re.split(r'\s+and\s+|\s*&\s*', sentence, flags=re.I)
        sentences = [re.sub(r'\s+', '', ctx) for ctx in sentences]
    else:
        sentences = re.split(r'\s+', sentence)
    # data["query"]["bool"]["must"] = [{"match_phrase": {"content": st}} for st in sentences]
    sentences = [re.split(r':|：', st) for st in sentences]
    # 标准化经号number字段,按照长度不同分别在number和sutra字段中查找
    must = (("content", st[0]) if len(st) == 1 else (st[0].lower(), st[1]) for st in sentences)
    # must = ((st0, st1 if st0 != 'number' else normalize_number(st1,False)) for st0,st1 in must)
    must = ((st0, st1 if st0 != 'number' else str(parse_number1(st1,False))) for st0,st1 in must)
    must = (('sutra' if (st0=='number' and '_' not in st1) else st0, st1) for st0,st1 in must)
    data["query"]["bool"]["must"] = [{"match_phrase" if key=="content" else "match": {key:val}} for key,val in must]
    # pprint.pprint(data["query"]["bool"]["must"])
    # 用于高亮的内容
    hlsentence = python_unescape(''.join([st[0] for st in sentences if len(st) == 1]))
    # e = time.time()
    # print(e-s)

    s = time.time()
    r= requests.get(url, json=data, timeout=10)
    result = r.json()
    # e = time.time()
    print(result)
    # {'error': {'root_cause': [{'type': 'index_not_found_exception', 'reason': 'no such index [cbeta]', 'resource.type': 'index_or_alias', 'resource.id': 'cbeta', 'index_uuid': '_na_', 'index': 'cbeta'}], 'type': 'index_not_found_exception', 'reason': 'no such index [cbeta]', 'resource.type': 'index_or_alias', 'resource.id': 'cbeta', 'index_uuid': '_na_', 'index': 'cbeta'}, 'status': 404}

    hits = result['hits']['hits']
    # value = result['hits']['total']['value']
    result = []
    for hit in hits:
        _source = hit["_source"]
        author = _source['author']
        juan = _source["number"].split('n')[0]
        hl = highlight(hlsentence, _source["raw"])
        # 文章内容高亮显示
        result.append({'hl': hl, 'an': f'/xml/{juan}/{_source["number"]}.xml#{hit["_id"]}',
                'title':_source['title'], 'author': author,
                'number': _source["number"]})

    result.sort(key=lambda x: pagerank(x['number']))

    return result


# with gzip.open('dict/cipin.json.gz') as fd:
#     cipind = json.load(fd)
#
# def zhuyin(txt, ruby=False, cipin=50):
#     '''對txt文本注音, ruby是否使用ruby語法'''
#     if not ruby:
#         content = ' '.join(lookinkangxi(i)['pinyin'].split(' ')[0] for i in txt)
#     else:
#         result = []
#         for zi in txt:
#             if cipind.get(zi, 0) < cipin:
#                 pinyin = lookinkangxi(zi)['pinyin'].split(' ')[0]
#                 if pinyin:
#                     zi = f"<ruby>{zi}<rt>{pinyin}</rt></ruby>"
#             result.append(zi)
#         content = ''.join(result)
#     return content


import difflib
from difflib import *

def diff_ctx(lctx, rctx):
    '''比较两个文本的不同, 使用html排版'''

    d = Differ()
    result = d.compare(lctx, rctx)

    lctx = []
    rctx = []
    for line in result:
        if line.startswith(' '):
            line = line[2:]
            lctx.append(line)
            rctx.append(line)
        elif line.startswith('- '):
            line = line[2:]
            lctx.append(f'<span class="red">{line}</span>')
        elif line.startswith('+ '):
            line = line[2:]
            rctx.append(f'<span class="red">{line}</span>')
        elif line.startswith('?'):
            continue
    lctx = ''.join(lctx)
    rctx = ''.join(rctx)
    # lctx = ''.join(f'<p>{line}</p>' for line in lctx.splitlines())
    # rctx = ''.join(f'<p>{line}</p>' for line in rctx.splitlines())

    return {'lfile': lctx, 'rfile': rctx}

def make_docx(ff, temp='', fanti=True):
    if not fanti:
        stc = STConvertor()
    et = ET.parse(ff)
    root = et.getroot()

    teiheader = root.findall("teiHeader")[0]
    title = ''.join(root.findall("teiHeader/fileDesc/titleStmt/title")[0].itertext())
    title = title.split('No.')[1].strip().split(maxsplit=1)[1]
    author = root.findall("teiHeader/fileDesc/titleStmt/author")
    # 应该用byline?
    if author:
        # author = author[0].text
        author = ''.join(author[0].itertext())
    if not author:
        author = ''

    fname = ff.split('/')[-1].split('_')[0]
    if fname == 'D11n8817':
        title = '佛說觀世音三昧經'
    if fname == 'T19n0920':
        title = '佛心經品亦通大隨求陀羅尼'

    # title, author, pid, ct, fname
    paras = root.findall("p")
    text = root.findall("text")[0]
    body = text.findall("body")[0]
    document = Document()
    head = document.add_heading(title, 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head = document.add_heading(author,1)
    head.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for p in body.iter():
        if p.tag == 'p':
            xmlid = p.attrib.get('{http://www.w3.org/XML/1998/namespace}id', None)

            ctx = normalize_text(''.join((normalize_space(t) for t in p.itertext())))
            if not fanti:
                ctx = stc.t2s(ctx)
            para = document.add_paragraph(ctx)
            paragraph_format = para.paragraph_format
            #print(dir(paragraph_format))
            paragraph_format.first_line_indent = Inches(0.25)

            # yield (xmlid, fname, author, title,  ctx)

        if p.tag == 'lg':
            xmlid = p.attrib['{http://www.w3.org/XML/1998/namespace}id']
            ctx = normalize_text(''.join((normalize_space(t) for t in p.itertext())))
            if not fanti:
                ctx = stc.t2s(ctx)
            para = document.add_paragraph(ctx, style='Intense Quote')
            #yield (xmlid, fname, author, title,  ctx)
    docxfname = ff.split('/')[-1][:-4]
    document.save(os.path.join(temp, f'{docxfname}.docx'))



def main():
    ''''''
    pass

def test():
    ''''''
    pass


if __name__ == "__main__":
    # main()
    # test()
    # print(get_all_juan('T02n0099'))
    # print(get_all_juan('GA031n0032'))
    # print(get_all_juan('J31nB269'))
    # print(get_all_juan('T19n0974A'))
    # print(get_next_page('T02n0099_001'))
    # print(get_next_page('T01n0002_001'))

    # ctx = '五<g ref="#CB22072">說</g>九種命終心三界<g ref="#CB29911">々</g><g ref="#CB29911">々</g>生各潤生心各有三故<note place="inline">已上'
    # print(rm_ditto_mark(ctx))
    #str_in = "a-kAra"
    #print(hk2iast(str_in))
    #print(hk2iastdeve(str_in))
    # print(convert2t('大佛顶'))
    # ss = Search()
    # for idx in ss.search('大佛頂'):
    #     print(idx)
    # # TODO:搜索t1000, t1000_001, T01n0001, T01n0001_001, T01n0001_p0001a01, T01,no.1,p.1a1
    #titlepatten = re.compile(r'([a-zA-Z][a-zA-Z]?)(\d\dn)?(\d\d\d\d)(_\d\d\d)?')
    #titlepatten.find('t1000')
    #pprint.pprint(mulu)
    # print(parse_number('CBETA, T14, no. 475, pp. 537c8-538a14'))
    # print(parse_number('CBETA 2019.Q2, Y25, no. 25, p. 411a5-7'))
    # print(parse_number('CBETA 2019.Q3, T20, no. 1113B, p. 498c12-17'))
    # print(parse_number('T20n1113B'))
    # print(parse_number('T20n1113'))
    # print(parse_number('T01n0001_p0001a01'))
    # print(parse_number('1113b'))
    # print(parse_number('1113'))
    # print(parse_number('100.3'))
    print(parse_number1('220').url)
    print(parse_number1('220').pages)
    # print(parse_number1('220.200'))
    # print(parse_number1('220.201'))
    # print(parse_number1('J32nB271'))
    # print(normalize_text('說</g>九種命終心三界'))
    #for i in fullsearch('止觀明靜'):
    #    print(i)
    # stc = STConvertor()
    # print(stc.t2s('那莫三𭦟多嚩日羅赦憾云〃哦'))
    # print(stc.s2t('安乐国'))
    # print(stc.detect('安乐国'))
    # print(stc.detect('那莫三𭦟多嚩日羅赦憾云〃哦'))
    # print(get_all_juan('T20n1113B'))
    # print(get_all_juan('T20n1113'))
    sentence = '非施者福 title:毘耶娑'
    sentence = '非施者福'
    # print(highlight(sentence, raw))
    #print(parse_number2('大正藏第九卷第七〇九页'))
    #print(parse_number2('大正藏第十九卷第16頁下'))

